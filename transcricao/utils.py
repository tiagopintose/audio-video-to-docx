import os
import subprocess
import tempfile
import json
import whisper
from docx import Document
from django.conf import settings
from datetime import datetime

# Variáveis globais
MODEL_DEVICE = os.getenv('MODEL_DEVICE', 'cpu')
MODEL = whisper.load_model("medium", device=MODEL_DEVICE)  # Load once


def _get_duration(caminho_wav: str) -> float:
    """
    Obtém a duração do ficheiro WAV em segundos usando ffprobe.
    """
    comando = [
        "ffprobe", "-v", "quiet", "-print_format", "json", "-show_format", caminho_wav
    ]
    resultado = subprocess.run(comando, capture_output=True, text=True, check=True)
    data = json.loads(resultado.stdout)
    return float(data["format"]["duration"])


def _split_audio(caminho_wav: str, segment_duration: int = 1800) -> list:
    """
    Divide o ficheiro WAV em segmentos de 'segment_duration' segundos.
    Retorna uma lista de caminhos para os segmentos.
    """
    duration = _get_duration(caminho_wav)
    segments = []
    start = 0
    while start < duration:
        end = min(start + segment_duration, duration)
        segment_path = tempfile.NamedTemporaryFile(suffix=".wav", delete=False).name
        comando = [
            "ffmpeg", "-y", "-i", caminho_wav,
            "-ss", str(start), "-t", str(end - start),
            "-ar", "16000", "-ac", "1", "-vn",
            segment_path
        ]
        subprocess.run(comando, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, check=True)
        segments.append(segment_path)
        start = end
    return segments


def _converter_para_wav(caminho_entrada: str) -> str:
    """
    Converte qualquer ficheiro de áudio/vídeo para WAV 16 kHz mono (sem compressão),
    garantindo máxima compatibilidade e precisão para o Whisper.
    """
    caminho_temp = tempfile.NamedTemporaryFile(suffix=".wav", delete=False)
    caminho_temp.close()

    comando = [
        "ffmpeg", "-y", "-i", caminho_entrada,
        "-ar", "16000", "-ac", "1", "-vn",
        caminho_temp.name
    ]
    subprocess.run(comando, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, check=True)
    return caminho_temp.name


def transcrever_audio(caminho_audio: str) -> str:
    """
    Transcreve um ficheiro de áudio ou vídeo em português e
    guarda o resultado em 'MEDIA_ROOT/transcricao.docx'.
    Divide em segmentos se o ficheiro for longo (>30 min).
    """
    try:
        # 🔹 Converter o áudio/vídeo para WAV 16 kHz mono
        caminho_wav = _converter_para_wav(caminho_audio)

        # 🔹 Verificar duração
        duration = _get_duration(caminho_wav)
        segment_duration = 900  # 15 minutos para reduzir uso de RAM

        if duration <= segment_duration:
            # Ficheiro curto: transcrever diretamente
            resultado = MODEL.transcribe(caminho_wav, language="portuguese")
            transcricao = resultado["text"].strip()
        else:
            # Ficheiro longo: dividir e transcrever segmentos
            segments = _split_audio(caminho_wav, segment_duration)
            transcricao_parts = []
            for i, segment in enumerate(segments):
                print(f"Transcrevendo segmento {i+1}/{len(segments)}...")
                resultado = MODEL.transcribe(segment, language="portuguese")
                transcricao_parts.append(resultado["text"].strip())
                # Limpar segmento
                os.remove(segment)
            transcricao = " ".join(transcricao_parts)

        # 🔹 Garantir que MEDIA_ROOT existe
        os.makedirs(settings.MEDIA_ROOT, exist_ok=True)
        
        # 🔹 Gerar nome único para o arquivo
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        arquivo_transcricao = os.path.join(settings.MEDIA_ROOT, f"transcricao_{timestamp}.docx")

        # 🔹 Criar novo documento DOCX
        doc = Document()
        doc.add_heading("Transcrição de Áudio/Vídeo", level=1)
        doc.add_heading(f"Data: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}", level=2)
        doc.add_paragraph(transcricao)
        
        # 🔹 Salvar documento
        try:
            doc.save(arquivo_transcricao)
            print(f"✅ Arquivo salvo em: {arquivo_transcricao}")
        except Exception as e:
            print(f"❌ Erro ao salvar arquivo: {e}")
            raise

        return "✅ Transcrição concluída com sucesso!"
    except Exception as e:
        return f"❌ Erro na transcrição: {e}"
    finally:
        # 🔹 Limpar o ficheiro temporário
        try:
            os.remove(caminho_wav)
        except OSError:
            pass
